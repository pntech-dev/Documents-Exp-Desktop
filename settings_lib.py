import os
import re
import shutil
import sqlite3 as sql

import config

from docx import Document
from datetime import datetime
from PyQt5 import QtWidgets, QtCore
from PyQt5.QtCore import QAbstractTableModel

class Tab:
    def progress_bar_update(self, progressBar, progress_step):
        """
        A function that updates the progress bar value\n
        progressBar - the transmitted progress bar, the value of which will be updated\n
        progress_step - the step for which the progress bar will be updated
        """
        progressBar.setValue(progressBar.value() + progress_step)

    def button_enabler(self, object, button):
        """
        The function enables and disables the passed button based on the passed object\n
        objects - The object based on which the button is enabled/disabled\n
        button - The button whose state is being controlled
        """
        if len(object) > 0:
            button.setEnabled(True)
        else:
            button.setEnabled(False)

    def get_folder(self, parent, lineEdit):
        """
        A function that returns the path of the selected folder and inserts it into the passed lineEdit\n
        lineEdit - The field where the path to the selected folder will be displayed
        """
        folder_path = QtWidgets.QFileDialog.getExistingDirectory(parent,"Выбрать папку",".")
        lineEdit.setText(folder_path)

    def get_file(self, parent, lineEdit):
        """
        A function that writes the path to the selected file to the passed lineEdit\n
        lineEdit - The string where the path to the selected file will be written
        """
        file_path, _ = QtWidgets.QFileDialog.getOpenFileName(parent, "Выберите файл", "", "Файлы Excel (*.docx *.doc)")
        lineEdit.setText(file_path)
    
    def get_groups(self, groups_path):
        """
        A function that returns a list of all groups
        groups_path - 
        """
        if os.path.exists(groups_path):
            groups = []
            for group in os.listdir(path=groups_path):
                groups.append(group)

            return groups
        else:
            return []
    
    def get_categories(self, group):
        """
        A function that returns a list of all categories in a group\n
        group - The transmitted group, the categories of which we receive
        """
        group_path = os.path.join(config.CURRENT_PATH_TO_GROUPS, group)
        categories = os.listdir(fr"{group_path}")

        categories = [i[0:-3] for i in categories]

        return categories
    
    def get_categories_from_folder(self, categories_path):
        """
        A function that returns a list of category names\n
        categories_path - The path to the selected folder with categories
        """
        categories = []
        for catigory in os.listdir(path=categories_path):
            categories.append(catigory[:catigory.rfind(".docx")])

        return categories

    def create_comboboxes(self, data, comboboxes):
        """
        A function that adds the transmitted data to the transmitted combobox\n
        data - List of group names\n
        comboboxes - List of comboboxes
        """
        for box in comboboxes:
            box.clear()

        for box in comboboxes:
            box.addItems(data)

    def create_categories_comboboxes(self, comboboxes):
        for box_1, box_2 in comboboxes.items():
            data = self.get_categories(group=box_1.currentText())
            self.create_comboboxes(data=data, comboboxes=[box_2])

    def create_table(self, tableView, headers, data, column_width):
        """
        Sets up and displays a table in the provided table view.\n
        tableView - The QTableView widget where the table will be displayed.\n
        headers - A list of column headers for the table.\n
        data - A list of data to populate the table.\n
        column_width - The width to set for the first column.
        """
        model = TableModel(headers=headers, data=data)
        tableView.setModel(model)

        tableView.setColumnWidth(0, column_width)
        tableView.horizontalHeader().setSectionResizeMode(1, QtWidgets.QHeaderView.Stretch)
    
    def get_table_view_data(self, group, categorie):
        """
        The function returns data to fill in the table\n
        group - The group in which the list will be deleted\n
        categorie - The category in which the list will be deleted
        """
        query = "SELECT name FROM sqlite_master WHERE type='table';"
        data = []

        if categorie:
            db_path = f"{os.path.join(config.CURRENT_PATH_TO_GROUPS, group, categorie)}.db"
            
            if os.path.exists(db_path):
                conn = sql.connect(database=db_path)
                cursor = conn.cursor()

                cursor.execute(query)
                db_data = cursor.fetchall()
                
                conn.commit()
                cursor.close()
                conn.close()

                sorted_tables = sorted(db_data, key=lambda name: self.extract_numbers(name[0].split("+")[0]))

                for name in sorted_tables:
                    data.append(name[0].split("+"))

        return data
    
    def extract_numbers(self, string):
        """
        A function that returns a tuple of numbers of all found numbers\n
        string - a string with numbers
        """
        numbers = re.findall(r'\d+', string)
        return tuple(map(int, numbers))
    
    def double_clicked(self, index, tableView, label, text):
        """
        The function changes the passed label if a double click was performed on a cell in the table\n
        index - cell in the table\n
        labele - the label that will change
        """
        row_data = []
        for col in range(tableView.model().columnCount()):
            row_data.append(index.sibling(index.row(), col).data())

        label.setText(f"{text} {row_data[0]}")

    def write_changelog(self, command):
        """
        A function that creates and/or writes to a CHANGELOG file\n
        command - The command that will be recorded
        """
        config.CURRENT_DIRECTORY_PATH = os.getcwd()
        with open(f"{config.CURRENT_DIRECTORY_PATH}\\CHANGELOG.txt", "a+", encoding="utf-8") as file:
            current_datetime = datetime.now()
            date = current_datetime.date().strftime("%d-%m-%Y")
            time = current_datetime.time().strftime("%H:%M:%S")

            command = f"{date} | {time} | {command}\n"
            file.write(command)

class AddTab(Tab):
    def __init__(self, tableView, progressBar):
        super().__init__()
        self.progressBar = progressBar
        tableView = tableView

    def add_group(self, filename):
        """
        A function that creates a group named filename\n
        filename - text from lineEdit
        """
        self.progressBar.setValue(config.MIN_PROGRESS)
        progress_step = int(config.MAX_PROGRESS / config.ADD_GROUP_PROGRESS_BAR_STEPS)

        group_file = f"{filename}"

        self.progress_bar_update(progressBar=self.progressBar, progress_step=progress_step)

        group_path = os.path.join(config.CURRENT_PATH_TO_GROUPS, group_file)

        self.progress_bar_update(progressBar=self.progressBar, progress_step=progress_step)

        os.makedirs(group_path, exist_ok=True)

        self.progressBar.setValue(config.MAX_PROGRESS)
        QtCore.QTimer.singleShot(1000, lambda: self.progressBar.setValue(config.MIN_PROGRESS))

        config.CHANGES = True
        self.write_changelog(command=f"CREATE_GROUP <{filename}>")

    def add_categories(self, group, categories_path):
        """
        A function that adds categories to the selected group.
        group - The group to which the categories will be added
        categories_path - The path to the folder with categories
        """
        path_to_group = os.path.join(config.CURRENT_PATH_TO_GROUPS, group)
        categories = self.get_categories_from_folder(categories_path=categories_path)

        progress_step = config.MAX_PROGRESS // len(categories) if categories else config.MAX_PROGRESS
        self.progressBar.setValue(config.MIN_PROGRESS)

        for index, category in enumerate(categories, start=1):
            db_name = f"{category}.db"
            db_path = os.path.join(path_to_group, db_name)

            docx_file_path = os.path.join(categories_path, f"{category}.docx")

            if not os.path.exists(docx_file_path) or not docx_file_path.endswith('.docx'):
                QtWidgets.QMessageBox.warning(None, "Ошибка", f"Файл {docx_file_path} не найден или не является документом .docx")
                continue

            try:
                doc = Document(docx=docx_file_path)

                if not doc.tables:
                    QtWidgets.QMessageBox.warning(None, "Ошибка", f"В файле {docx_file_path} отсутствуют таблицы.")
                    continue

                tables_rows = []
                invalid_format = False

                for table in doc.tables:
                    for row in table.rows:
                        row_data = [cell.text.strip() for cell in row.cells]
                        if len(row_data) > 2:
                            invalid_format = True
                            break
                    if invalid_format:
                        break

                if invalid_format:
                    QtWidgets.QMessageBox.warning(None, "Неверный формат данных", 
                                                f"Не правильный формат данных в файле {docx_file_path}")
                    continue

                if not os.path.exists(db_path):
                    with open(db_path, 'w'):
                        pass

                    conn = sql.connect(db_path)
                    cursor = conn.cursor()

                    for table in doc.tables:
                        for row in table.rows:
                            row_data = [cell.text.strip() for cell in row.cells]
                            if len(row_data) == 2 and any(row_data):
                                table_name = f"{row_data[0].replace(' ', '')}+{row_data[1]}"
                                query_create = f"""
                                CREATE TABLE IF NOT EXISTS [{table_name}] (
                                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                                    Обозначение TEXT,
                                    Наименование TEXT
                                );
                                """
                                cursor.execute(query_create)

                    conn.commit()
                    cursor.close()
                    conn.close()

                    QtWidgets.QMessageBox.information(None, "Категория добавлена", f"Категория: {category} добавлена")
                    config.CHANGES = True
                    self.write_changelog(command=f"CREATE_CATEGORY <{group}> <{category}>")

                else:
                    doc = Document(docx=docx_file_path)
                    expected_tables = {
                        f"{row.cells[0].text.strip().replace(' ', '')}+{row.cells[1].text.strip()}"
                        for table in doc.tables for row in table.rows if len(row.cells) == 2 and any(cell.text.strip() for cell in row.cells)
                    }

                    conn = sql.connect(db_path)
                    cursor = conn.cursor()
                    cursor.execute("SELECT name FROM sqlite_master WHERE type='table';")
                    existing_tables = {table[0] for table in cursor.fetchall()}
                    conn.commit()
                    cursor.close()
                    conn.close()

                    missing_tables = expected_tables - existing_tables

                    if missing_tables:
                        conn = sql.connect(db_path)
                        cursor = conn.cursor()
                        for table_name in missing_tables:
                            query_create = f"""
                            CREATE TABLE IF NOT EXISTS [{table_name}] (
                                id INTEGER PRIMARY KEY AUTOINCREMENT,
                                Обозначение TEXT,
                                Наименование TEXT
                            );
                            """
                            cursor.execute(query_create)
                        conn.commit()
                        cursor.close()
                        conn.close()

                        QtWidgets.QMessageBox.information(None, "Данные обновлены", f"Категория: {category} обновлена")
                        config.CHANGES = True
                        self.write_changelog(command=f"CHANGES_IN_CATEGORY <{group}> <{category}>")
                    elif expected_tables == existing_tables:
                        QtWidgets.QMessageBox.information(None, "Категория существует", f"Категория: {category} уже существует")

                self.progressBar.setValue(index * progress_step)

            except Exception as e:
                QtWidgets.QMessageBox.warning(None, "Ошибка", f"Не удалось обработать категорию {category}: {str(e)}")
                continue

        QtCore.QTimer.singleShot(1000, lambda: self.progressBar.setValue(config.MIN_PROGRESS))

    def add_lists(self, group, categorie, file_path):
        """
        A function that adds lists to the transferred category in the transferred group.
        group - The group to which the lists will be added to.
        categorie - The category to which the lists will be added.
        file_path - The path to the file with the lists.
        """
        self.progressBar.setValue(config.MIN_PROGRESS)

        try:
            doc = Document(docx=file_path)
        except Exception as e:
            QtWidgets.QMessageBox.warning(None, "Ошибка", f"Не удалось открыть файл {file_path}: {str(e)}")
            return

        doc_data = {}
        for table in doc.tables:
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                if len(row_data) == 3 and any(row_data):
                    table_name = row_data[2]
                    if table_name not in doc_data:
                        doc_data[table_name] = []
                    doc_data[table_name].append((row_data[0], row_data[1]))

        if not doc_data:
            QtWidgets.QMessageBox.warning(None, "Неверный формат данных", f"Не правильный формат данных в файле {file_path}")
            return

        categorie_path = f"{os.path.join(config.CURRENT_PATH_TO_GROUPS, group, categorie)}.db"
        if not os.path.exists(categorie_path):
            QtWidgets.QMessageBox.warning(None, "Ошибка", f"База данных для категории {categorie} не найдена.")
            return

        conn = sql.connect(database=categorie_path)
        cursor = conn.cursor()

        query = "SELECT name FROM sqlite_master WHERE type='table';"
        cursor.execute(query)
        db_tables = {row[0].split('+')[0]: row[0] for row in cursor.fetchall()}

        progress_step = config.MAX_PROGRESS // (len(doc_data) + len(db_tables))

        for doc_table_name, doc_rows in doc_data.items():
            if doc_table_name in db_tables:
                full_table_name = db_tables[doc_table_name]

                # Check if id column exists
                cursor.execute(f"PRAGMA table_info([{full_table_name}])")
                columns = [info[1] for info in cursor.fetchall()]
                if 'id' not in columns:
                    # Migrate table
                    cursor.execute(f'SELECT * FROM [{full_table_name}]')
                    old_data = cursor.fetchall()
                    
                    cursor.execute(f'DROP TABLE [{full_table_name}]')

                    query_create = f"""
                    CREATE TABLE IF NOT EXISTS [{full_table_name}] (
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        Обозначение TEXT,
                        Наименование TEXT
                    );
                    """
                    cursor.execute(query_create)

                    cursor.executemany(f'INSERT INTO [{full_table_name}] (Обозначение, Наименование) VALUES (?, ?)', old_data)

                query = f'SELECT "Обозначение", "Наименование" FROM [{full_table_name}]'
                cursor.execute(query)
                db_rows = set(cursor.fetchall())

                rows_to_add = [row for row in doc_rows if row not in db_rows]

                if rows_to_add:
                    query = f'INSERT INTO [{full_table_name}] ("Обозначение", "Наименование") VALUES (?, ?)'
                    cursor.executemany(query, rows_to_add)
            
            self.progress_bar_update(progressBar=self.progressBar, progress_step=progress_step)

        conn.commit()
        cursor.close()
        conn.close()

        self.progressBar.setValue(config.MAX_PROGRESS)
        QtCore.QTimer.singleShot(1000, lambda: self.progressBar.setValue(config.MIN_PROGRESS))

        config.CHANGES = True
        self.write_changelog(command=f"CHANGES_IN_CATEGORY <{group}> <{categorie}>")

class DeleteTab(Tab):
    def __init__(self, tableView, progressBar):
        super().__init__()
        self.progressBar = progressBar
        self.tableView = tableView

    def delete_group(self, group):
        """
        The function deletes the passed group\n
        group - The group to be deleted
        """
        self.progressBar.setValue(config.MIN_PROGRESS)
        progress_step = int(config.MAX_PROGRESS / config.DELETE_GROUP_PROGRESS_BAR_STEPS)

        group_path = os.path.join(config.CURRENT_PATH_TO_GROUPS, group)
        self.progress_bar_update(progressBar=self.progressBar, progress_step=progress_step)

        shutil.rmtree(group_path)

        self.progressBar.setValue(config.MAX_PROGRESS)
        QtCore.QTimer.singleShot(1000, lambda: self.progressBar.setValue(config.MIN_PROGRESS))

        config.CHANGES = True
        self.write_changelog(command=f"DELETE_GROUP <{group}>")

    def delete_categorie(self, group, categorie):
        """
        The function deletes the passed group\n
        group - The group in which the category will be deleted\n
        categorie - the category to be deleted
        """
        self.progressBar.setValue(config.MIN_PROGRESS)
        progress_step = int(config.MAX_PROGRESS / config.DELETE_CATEGORIE_PROGRESS_BAR_STEPS)

        categorie_path = f"{os.path.join(config.CURRENT_PATH_TO_GROUPS, group, categorie)}.db"
        self.progress_bar_update(progressBar=self.progressBar, progress_step=progress_step)

        os.remove(path=categorie_path)

        self.progressBar.setValue(config.MAX_PROGRESS)
        QtCore.QTimer.singleShot(1000, lambda: self.progressBar.setValue(config.MIN_PROGRESS))

        config.CHANGES = True
        self.write_changelog(command=f"DELETE_CATEGORY <{group}> <{categorie}>")

    def delete_list(self, group, categorie, label):
        """
        The function removes the list recorded in the label from the passed group and category\n
        group - The group in which the list will be deleted\n
        categorie - The category in which the list will be deleted\n
        label - The name of the list to be deleted
        """
        self.progressBar.setValue(config.MIN_PROGRESS)
        progress_step = int(config.MAX_PROGRESS / config.DELETE_LIST_PROGRESS_BAR_STEPS)

        categorie_path = f"{os.path.join(config.CURRENT_PATH_TO_GROUPS, group, categorie)}.db"

        self.progress_bar_update(progressBar=self.progressBar, progress_step=progress_step)
        
        if os.path.exists(categorie_path):
            conn = sql.connect(database=categorie_path)
            cursor = conn.cursor()

            self.progress_bar_update(progressBar=self.progressBar, progress_step=progress_step)

            query = "SELECT name FROM sqlite_master WHERE type='table';"
            cursor.execute(query)
            tables = cursor.fetchall()

            self.progress_bar_update(progressBar=self.progressBar, progress_step=progress_step)
    
            for table in tables:
                if table[0].split("+")[0] == label[17:]:
                    query = f"""DROP TABLE IF EXISTS [{table[0]}]"""
                    cursor.execute(query)

            self.progress_bar_update(progressBar=self.progressBar, progress_step=progress_step)

            conn.commit()
            cursor.close()
            conn.close()

        self.progressBar.setValue(config.MAX_PROGRESS)
        QtCore.QTimer.singleShot(1000, lambda: self.progressBar.setValue(config.MIN_PROGRESS))

        config.CHANGES = True
        self.write_changelog(command=f"CHANGES_IN_CATEGORY <{group}> <{categorie}>")

    def search(self, tableView, search_text, group, categorie):
        """
        A function that searches for a given value in a given category\n
        group - The group in which the search will be performed\n
        category - The category in which the search will be performed
        """

        if len(search_text) > 0:
            db_path = f"{os.path.join(config.CURRENT_PATH_TO_GROUPS, group, categorie)}.db"

            conn = sql.connect(database=db_path)
            cursor = conn.cursor()

            data = []
            search_text = search_text.lower().strip()

            query = "SELECT name FROM sqlite_master WHERE type='table';"
            cursor.execute(query)
            tables = cursor.fetchall()

            for table in tables:
                table_name = table[0].lower().strip()
                if all(word in table_name for word in search_text.split()):
                    data.append(table[0].split("+"))

                table_name = table[0]
                query = f"SELECT * FROM [{table_name}]"
                cursor.execute(query)
                table_data = cursor.fetchall()

                for lst in table_data:
                    lst = [i.replace("\n", " ") for i in lst]
                    row = " ".join(lst)
                    
                    if all(word in row.lower().strip() for word in search_text.split()):
                        data.append([table[0].split("+")[0], row])

            self.create_table(tableView=tableView, headers=["№", "Наименование"], data=data, column_width=70)

            conn.commit()
            cursor.close()
            conn.close()
        else:
            self.create_table(tableView=tableView, headers=["№", "Наименование"], data=self.get_table_view_data(group=group, categorie=categorie), column_width=70)

class TemplatesTab(Tab):
    def __init__(self, tableView, progressBar):
        super().__init__()
        self.tableView = tableView
        self.progressBar = progressBar

    def create_templates_tables(self, tableView, headers):
        """
        The function creates a table with template names\n
        tableView - the table that will be filled in\n
        headers - column header
        """
        templates = []
        templates_names = os.listdir(config.PATH_TO_TEMPLATES)
        for template in templates_names:
            if ".docx" in template:
                templates.append([template])

        model = TableModel(headers=headers, data=templates)
        tableView.setModel(model)

        tableView.horizontalHeader().setSectionResizeMode(0, QtWidgets.QHeaderView.Stretch)

    def download(self, line_edit, file_label):
        """
        The function downloads the selected file from the group to the specified directory\n
        line_edit - the path to save\n
        group_label - the place where the selected group is indicated\n
        file_label - The location where the selected file is specified
        """
        self.progressBar.setValue(config.MIN_PROGRESS)
        progress_step = int(config.MAX_PROGRESS / config.TEMPLATES_PROGRESS_BAR_STEPS)

        file = file_label[15:].strip()
        template_path = os.path.join(config.PATH_TO_TEMPLATES, file)

        if not os.path.isfile(template_path):
            QtWidgets.QMessageBox.warning(None, "Ошибка", f"Файл шаблона не найден по пути: {template_path}")
            return

        if not line_edit:
            save_path = os.path.join(os.path.expanduser('~'), 'Desktop', file)
        else:
            save_path = os.path.join(line_edit, file)

        self.progress_bar_update(progressBar=self.progressBar, progress_step=progress_step)

        try:
            shutil.copy(template_path, save_path)
        except Exception as e:
            QtWidgets.QMessageBox.warning(None, "Ошибка", f"Не удалось скопировать файл: {e}")
            self.progressBar.setValue(config.MIN_PROGRESS)
            return

        self.progressBar.setValue(config.MAX_PROGRESS)
        QtCore.QTimer.singleShot(1000, lambda: self.progressBar.setValue(config.MIN_PROGRESS))

class TableModel(QAbstractTableModel):
    def __init__(self, headers, data, parent = None):
        super().__init__(parent)
        
        self.headers = headers
        self._data = data

    def rowCount(self, index):
        """
        Returns the number of rows in the model\n
        index - The index of the item in the model (not used in this method)
        """
        return len(self._data)

    def columnCount(self, index=None):
        """
        Returns the number of columns in the model\n
        index - The index of the item in the model (not used in this method)
        """
        return len(self.headers)
    
    def data(self, index, role):
        """
        Returns the data for a specific item in the model\n
        index - The index of the item (row and column) in the model\n
        role - The role of the data (e.g., DisplayRole, TextWordWrap)
        """
        if role == QtCore.Qt.DisplayRole:
            try:
                return self._data[index.row()][index.column()]
            except IndexError:
                return None
        
    def headerData(self, section, orientation, role):
        """
        Returns the header data for a specified section of the table\n
        section - The section (column) index\n
        orientation - The orientation of the header (horizontal or vertical)\n
        role - The role of the data
        """
        if role == QtCore.Qt.DisplayRole:
            if orientation == QtCore.Qt.Horizontal:
                return self.headers[section]