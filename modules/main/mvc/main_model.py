import docx
import keyring
import logging
import requests

from pathlib import Path

from api.api_client import APIClient
from utils.app_paths import get_app_data_dir, get_local_data_dir
from utils.file_utils import load_config, read_json


class MainModel:
    def __init__(self) -> None:
        self.config_data = load_config()

        # API Client Initialization
        self.api = APIClient(base_url=self.config_data.get("base_url", None))
        
        # Constants
        self.APP_DIR = get_app_data_dir()
        self.LOCAL_DIR = get_local_data_dir()
        self.LOCAL_DIR_LAST_LOGGED = self.LOCAL_DIR / "last_logged.json"

        # Sidebar data
        self.departments = self._get_departments()
        self.current_department_id = self.departments[0]["id"] if self.departments else None
        self.categories = self._get_categories()
        self.current_category_id = self.categories[0]["id"] if self.categories else None

        # Table data
        self.documents = [] # Will be loaded via pagination
        # (x, y): x - id, bool(True, False) if document is page
        self.selected_document = (None, False)



    def get_user_data(self) -> dict | None:
        token = self._get_user_token()

        if not token:
            return None
        
        data = self.api.get_user_data(token)
        
        return data
    

    def get_document(self, document_id: int) -> dict:
        """Retrieves a specific document by ID."""
        return self.api.get_document(document_id)


    def get_document_pages(
            self, 
            document_id: int
    ) -> list[dict]:
        pages = self.api.get_document_pages(document_id)
        return pages["pages"]
    

    def create_department(self, name: str) -> dict:
        """Create new department"""
        data = {"name": name}
        return self._make_authorized_request(self.api.create_department, data=data)


    def edit_department(self, name: str, department_id: int):
        """Edit department data"""
        data = {"name": name}
        return self._make_authorized_request(self.api.edit_department, data=data, department_id=department_id)


    def delete_department(self, department_id: int):
        """Delete department"""
        return self._make_authorized_request(self.api.delete_department, department_id=department_id)
    

    def create_category(self, name: str) -> dict:
        """Create new category"""
        data = {
            "group_id": self.current_department_id,
            "name": name
        }
        return self._make_authorized_request(self.api.create_category, data=data)


    def refresh_data(self) -> None:
        """Refreshes the data from the API."""
        self.departments = self._get_departments()
        self.categories = self._get_categories()
        # Documents will be refreshed by the controller resetting pagination


    def fetch_documents(self, category_id: int, limit: int, offset: int) -> list[dict]:
        """Fetches a page of documents for a specific category."""
        response = self.api.get_documents(category_id=category_id, limit=limit, offset=offset)
        return response.get("documents", [])


    def search_data(self, query: str) -> list[dict]:
        """Searches the data based on the provided query."""
        results = []

        if not query:
            return results

        results = self.api.search_data(
            category_id=self.current_category_id, 
            query=query
        )

        return results


    def export_to_docx(
            self, 
            path: str, 
            filename: str, 
            data: dict = None
    ) -> None:
        """Exports the document data to a Word document.

        Args:
            path: The path to save the Word document.
            filename: The name of the Word document.
            data: The document data as a dictionary.
        """
        doc = docx.Document()

        doc.add_heading(f"{data.get('department')} - {data.get('category')}", 1)

        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'

        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Код'
        hdr_cells[1].text = 'Наименование'

        for document in data.get("documents", []):
            row_cells = table.add_row().cells
            row_cells[0].text = str(document.get("code", "") or "")
            row_cells[1].text = str(document.get("name", "") or "")

        full_path = Path(path) / filename
        doc.save(str(full_path))


    # ====================
    # Model Methods
    # ====================

    def _make_authorized_request(self, api_method, **kwargs):
        """Executes an API request with automatic token refresh on 401."""
        try:
            token = self._get_user_token()
            return api_method(token=token, **kwargs)
        except requests.exceptions.HTTPError as e:
            if e.response is not None and e.response.status_code == 401:
                logging.info("Access token expired. Refreshing...")
                if self._refresh_tokens():
                    token = self._get_user_token()
                    return api_method(token=token, **kwargs)
            raise e

    def _refresh_tokens(self) -> bool:
        """Refreshes the access token using the refresh token."""
        try:
            last_logged = read_json(self.LOCAL_DIR_LAST_LOGGED)
            if not last_logged:
                return False
            
            user_id = last_logged.get("user_id")
            refresh_token = keyring.get_password("Documents Exp", f"refresh_token_{user_id}")
            
            if not refresh_token:
                return False
                
            tokens = self.api.refresh(refresh_token)
            keyring.set_password("Documents Exp", f"access_token_{user_id}", tokens["access_token"])
            keyring.set_password("Documents Exp", f"refresh_token_{user_id}", tokens["refresh_token"])
            return True
        except Exception as e:
            logging.error(f"Failed to refresh tokens: {e}")
            return False

        

    def _get_user_token(self) -> str | None:
        last_logged = read_json(self.LOCAL_DIR_LAST_LOGGED)

        if not last_logged:
            return None

        user_id = last_logged["user_id"]

        access_token = keyring.get_password(
            service_name="Documents Exp",
            username=f"access_token_{user_id}"
        )

        return access_token
    

    def _get_departments(self) -> list[dict]:
        departments = self.api.get_departments()
        return departments["departments"]


    def _get_categories(self) -> list[dict]:
        categories = self.api.get_categories()
        return categories["categories"]
    