import docx
import logging
import keyring
import requests
import os

from pathlib import Path

from api.api_client import APIClient
from utils.app_paths import get_app_data_dir, get_local_data_dir
from utils.file_utils import load_config, read_json



class DocumentEditorModel:
    def __init__(
            self,
            category_id: int = None,
            document_data: dict = None, 
            pages: list[dict] = None
    ) -> None:
        self.category_id = category_id
        self.document_data = document_data if document_data is not None else {}
        self.pages = pages if pages is not None else []
        self.pending_files = []
        self._pending_file_keys = set()

        self.config_data = load_config()

        # Constants
        self.APP_DIR = get_app_data_dir()
        self.LOCAL_DIR = get_local_data_dir()
        self.LOCAL_DIR_LAST_LOGGED = self.LOCAL_DIR / "last_logged.json"

        # API Client Initialization
        self.api = APIClient(base_url=self.config_data.get("base_url", None))

        self.is_document_edited = False


    def import_from_docx(self, file_path: str) -> list[dict]:
        """Imports pages from a Word document."""
        doc = docx.Document(file_path)
        pages = []

        for table in doc.tables:
            if not table.rows:
                continue
            
            # Check headers in the first row
            headers = [cell.text.strip() for cell in table.rows[0].cells]
            if len(headers) != 2:
                continue
            
            # Flexible header search
            headers_map = {h.lower(): i for i, h in enumerate(headers)}
            
            name_synonyms = ["наименование", "название", "имя", "name", "title"]
            code_synonyms = ["код", "шифр", "обозначение", "code", "designation"]
            
            name_index = next((headers_map[s] for s in name_synonyms if s in headers_map), -1)
            code_index = next((headers_map[s] for s in code_synonyms if s in headers_map), -1)

            if name_index != -1 and code_index != -1:

                # Iterate over data rows
                for row in table.rows[1:]:
                    cells = row.cells
                    if len(cells) > max(name_index, code_index):
                        name_text = cells[name_index].text.strip()
                        code_text = cells[code_index].text.strip()
                        pages.append({
                            "id": None, # New page has no ID
                            "name": name_text,
                            "designation": code_text
                        })
                
                # Stop after finding and processing the first valid table
                if pages:
                    break
        
        return pages
    

    def export_to_docx(
            self, 
            path: str, 
            filename: str, 
            data: dict = None
    ) -> None:
        """Exports the document data to a Word document.

        Args:
            path: The path to save the Word document.
            filename: The name of the Word document.
            data: The document data as a dictionary.
        """
        doc = docx.Document()

        if data:
            code = data.get("code", "")
            name = data.get("name", "")
            pages = data.get("pages", [])
        else:
            code = self.document_data.get("code", "")
            name = self.document_data.get("name", "")
            pages = self.pages
            pages.sort(key=lambda x: x.get("order_index", 0))

        doc.add_heading(f"{code} - {name}", 1)

        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'

        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Наименование'
        hdr_cells[1].text = 'Код'

        for page in pages:
            row_cells = table.add_row().cells
            row_cells[0].text = str(page.get("name", "") or "")
            row_cells[1].text = str(page.get("designation", "") or "")

        full_path = Path(path) / filename
        doc.save(str(full_path))


    def download_file(self, file_id: int, save_path: str) -> None:
        """Downloads a file."""
        self._make_authorized_request(self.api.download_file, file_id=file_id, destination_path=save_path)


    def upload_file(self, file_path: str) -> dict:
        """Uploads a file to the current document."""
        doc_id = self.document_data.get("id")
        if doc_id is None:
            raise ValueError("Сначала сохраните документ перед загрузкой файлов.")
            
        return self._make_authorized_request(self.api.upload_file, document_id=doc_id, file_path=file_path)


    @staticmethod
    def _pending_file_key(file_path: str) -> str:
        """Normalizes file path for duplicate detection."""
        try:
            normalized = str(Path(file_path).expanduser().resolve(strict=False))
        except Exception:
            normalized = str(file_path)
        return normalized.lower() if os.name == "nt" else normalized

    def _rebuild_pending_file_keys(self) -> None:
        self._pending_file_keys = {self._pending_file_key(path) for path in self.pending_files}

    def add_pending_file(self, file_path: str) -> bool:
        """Adds a file to the pending upload list.

        Returns:
            bool: True if file added, False if duplicate.
        """
        self._rebuild_pending_file_keys()
        key = self._pending_file_key(file_path)
        if key in self._pending_file_keys:
            return False
        self.pending_files.append(file_path)
        self._pending_file_keys.add(key)
        return True

    def remove_pending_file(self, file_path: str) -> None:
        """Removes a file from the pending upload list."""
        self._rebuild_pending_file_keys()
        remove_key = self._pending_file_key(file_path)
        to_remove = None
        for path in self.pending_files:
            if self._pending_file_key(path) == remove_key:
                to_remove = path
                break
        if to_remove is not None:
            self.pending_files.remove(to_remove)
            self._pending_file_keys.discard(remove_key)


    def upload_pending_files(self) -> None:
        """Uploads all pending files."""
        self._rebuild_pending_file_keys()
        errors = []
        for file_path in list(self.pending_files):
            try:
                self.upload_file(file_path)
                self.pending_files.remove(file_path)
                self._pending_file_keys.discard(self._pending_file_key(file_path))
            except Exception as e:
                errors.append(f"Не удалось загрузить {Path(file_path).name}: {e}")
        
        if errors:
            raise Exception("\n".join(errors))

    def delete_file(self, file_id: int) -> None:
        """Deletes a file from the document."""
        self._make_authorized_request(self.api.delete_file, file_id=file_id)


    def delete_document(self) -> None:
        self._make_authorized_request(
            self.api.delete_document,
            document_id=self.document_data.get("id")
        )


    def save_document(self, data: dict) -> None:
        doc_id = self.document_data.get("id")
        if doc_id is None:
            data["category_id"] = self.category_id
            response = self._make_authorized_request(self.api.create_document, data=data)
            self.document_data = response
        else:
            self._make_authorized_request(self.api.update_document, document_id=doc_id, data=data)


    # ====================
    # Model Methods
    # ====================

    
    def _get_user_token(self) -> str | None:
        last_logged = read_json(self.LOCAL_DIR_LAST_LOGGED)
        if not isinstance(last_logged, dict):
            return None

        user_id = last_logged.get("user_id")
        if user_id in (None, ""):
            return None

        access_token = keyring.get_password(
            service_name="Documents Exp",
            username=f"access_token_{user_id}"
        )

        return access_token


    def _make_authorized_request(self, api_method, **kwargs):
        """Executes an API request with automatic token refresh on 401."""
        try:
            token = self._get_user_token()
            return api_method(token=token, **kwargs)
        except requests.exceptions.HTTPError as e:
            if e.response is not None and e.response.status_code == 401:
                logging.info("Access token expired. Refreshing...")
                if self._refresh_tokens():
                    token = self._get_user_token()
                    return api_method(token=token, **kwargs)
            raise e


    def _refresh_tokens(self) -> bool:
        """Refreshes tokens using the stored refresh token."""
        try:
            last_logged = read_json(self.LOCAL_DIR_LAST_LOGGED)
            if not isinstance(last_logged, dict):
                return False
            
            user_id = last_logged.get("user_id")
            if user_id in (None, ""):
                return False
            refresh_token = keyring.get_password("Documents Exp", f"refresh_token_{user_id}")
            
            if not refresh_token:
                return False
                
            tokens = self.api.refresh(refresh_token)
            keyring.set_password("Documents Exp", f"access_token_{user_id}", tokens["access_token"])
            keyring.set_password("Documents Exp", f"refresh_token_{user_id}", tokens["refresh_token"])
            return True
        except Exception as e:
            logging.error(f"Failed to refresh tokens: {e}")
            return False
