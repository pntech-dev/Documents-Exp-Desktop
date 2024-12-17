import sqlite3 as sql

from docx import Document

from config import *
from PyQt5 import QtWidgets, QtCore, QtGui

def button_clicked(lineEdit, document_path, table_view):
    lineEdit.setEnabled(True)

    doc = Document(document_path)

    all_data = []

    if not lineEdit.text():
        for table in doc.tables:
            for row in table.rows:
                all_data.append([cell.text for cell in row.cells])

    else:
        lineEdit_words = lineEdit.text().lower().split()
        
        for table in doc.tables:
            for row in table.rows:
                row_text = " ".join([cell.text.lower() for cell in row.cells])
                
                if all(word in row_text for word in lineEdit_words):
                    all_data.append([cell.text for cell in row.cells])
        
                with sql.connect(LISTS) as conn:
                    cursor = conn.cursor()

                    i = [cell.text.lower() for cell in row.cells][0].upper().replace(" ", "").replace("/", "-")
                    i = i.upper().replace(" ", "").replace("/", "-")
                    
                    try:
                        cursor.execute(f"SELECT * FROM [{i}]")
                        rows = cursor.fetchall()

                        for row in rows:
                            table_row_text = " ".join([table_word.lower() for table_word in row])
                                
                            if all(word in table_row_text for word in lineEdit_words):
                                all_data.append([i.replace("-", "/"), f"{row[0]} {row[1]}"])
                    except:
                        print(f"{i} - Not in database")

    if all_data:
        create_table(data=all_data, headers=["№", "Наименование"], table_view=table_view, column_width=60)
    else:
        QtWidgets.QMessageBox.warning(None, "Ошибка", "Данные не найдены.")

def double_clicked(table, index, lineEdit, label):
    model = table.model()
    row = index.row()
    row_data = [model.data(model.index(row, col), QtCore.Qt.DisplayRole) for col in range(model.columnCount())]
    row_data_0 = row_data[0].replace(" ", "").replace("/", "-")

    all_data = []
    with sql.connect(LISTS) as conn:
        cursor = conn.cursor()

        cursor.execute("SELECT name FROM sqlite_master WHERE type='table';")
        tables = [table[0] for table in cursor.fetchall()]

        if row_data_0 in tables:
            lineEdit.setEnabled(False)
            
            cursor.execute(f"SELECT * FROM [{row_data_0}]")
            rows = cursor.fetchall()

            label.setText(f"{label.text()[0:9]} {row_data[0]}")
            
            for row in rows:
                all_data.append(row)
            
            create_table(data=all_data, headers=["Обозначение", "Наименование"], table_view=table, column_width=200)

        conn.commit()

def create_table(data, headers, table_view, column_width):
    model = TableModel(data, headers)
    table_view.setModel(model)

    table_view.setColumnWidth(0, column_width)
    table_view.horizontalHeader().setSectionResizeMode(1, QtWidgets.QHeaderView.Stretch)

    header_font = QtGui.QFont("Segoe", 12)
    table_view.horizontalHeader().setFont(header_font)

    table_view.setWordWrap(True)
    table_view.resizeRowsToContents()

class TableModel(QtCore.QAbstractTableModel):
    def __init__(self, data, headers):
        super(TableModel, self).__init__()
        self._data = data
        self.headers = headers

    def data(self, index, role):
        if role == QtCore.Qt.DisplayRole:
            return self._data[index.row()][index.column()]

        if role == QtCore.Qt.TextWordWrap:
            return True

    def rowCount(self, index):
        return len(self._data)

    def columnCount(self, index=None):
        return len(self.headers)

    def headerData(self, section, orientation, role):
        if role == QtCore.Qt.DisplayRole:
            if orientation == QtCore.Qt.Horizontal:
                return self.headers[section]